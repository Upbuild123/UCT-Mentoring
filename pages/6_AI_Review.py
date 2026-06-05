import sys
sys.path.insert(0, ".")
from io import BytesIO
import streamlit as st
import db

st.set_page_config(page_title="AI Coaching Review", layout="centered")
st.markdown("<style>[data-testid='stSidebarNav'] {display: none;}</style>", unsafe_allow_html=True)

# Reduce heading sizes to match body text — professional document style
st.markdown("""
<style>
h1 { font-size: 1.05rem !important; font-weight: 700 !important; margin: 1rem 0 0.2rem 0 !important; }
h2 { font-size: 1.0rem !important; font-weight: 700 !important; margin: 0.8rem 0 0.2rem 0 !important; }
h3 { font-size: 0.97rem !important; font-weight: 700 !important; margin: 0.6rem 0 0.2rem 0 !important; }
</style>
""", unsafe_allow_html=True)

st.title("AI Coaching Review")

params = st.query_params
assessment_id_str = params.get("assessment_id")

if not assessment_id_str:
    st.error("No assessment ID provided.")
    st.stop()

try:
    assessment_id = int(assessment_id_str)
except ValueError:
    st.error("Invalid assessment ID.")
    st.stop()

assessment = db.get_assessment_by_id(assessment_id)
if not assessment:
    st.error("Assessment not found.")
    st.stop()

student = db.get_student_by_id(assessment["student_id"])
base_name = f"Mentoring Round {assessment['round']}. {student['name']}. AI Review"
st.caption(f"{student['name']} -- Round {assessment['round']}")
st.divider()

ai_review = db.get_ai_review(assessment_id)
if not ai_review:
    st.info("AI review not yet available.")
    st.stop()

content = ai_review["content"]
st.markdown(content)

st.divider()

def to_docx(text: str, title: str) -> bytes:
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    doc.add_heading(title, 0)
    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue
        if stripped.startswith("**") and stripped.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip("*"))
            run.bold = True
        elif stripped.startswith("*") and stripped.endswith("*"):
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip("*"))
            run.italic = True
        else:
            doc.add_paragraph(stripped)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()

col1, col2 = st.columns(2)
with col1:
    st.download_button(
        label="Download as .txt",
        data=content.encode("utf-8"),
        file_name=f"{base_name}.txt",
        mime="text/plain",
    )
with col2:
    st.download_button(
        label="Download as .docx",
        data=to_docx(content, base_name),
        file_name=f"{base_name}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
