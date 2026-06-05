import sys
sys.path.insert(0, ".")
from io import BytesIO
import streamlit as st
import db

st.title("Session Transcript")

params = st.query_params
assessment_id_str = params.get("assessment_id")

if not assessment_id_str:
    st.error("No assessment ID provided.")
    st.stop()

try:
    assessment_id = int(assessment_id_str)
except ValueError:
    st.error("Invalid assessment ID.")
    st.stop()

assessment = db.get_assessment_by_id(assessment_id)
if not assessment:
    st.error("Assessment not found.")
    st.stop()

student = db.get_student_by_id(assessment["student_id"])
base_name = f"Mentoring Round {assessment['round']}. {student['name']}. Transcript"
st.caption(f"{student['name']} -- Round {assessment['round']}")
st.divider()

transcript = assessment.get("transcript") or "(Transcript not yet available.)"
st.markdown(transcript)

st.divider()

def to_docx(text: str, title: str) -> bytes:
    from docx import Document
    doc = Document()
    doc.add_heading(title, 0)
    for line in text.split("\n"):
        doc.add_paragraph(line)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()

col1, col2 = st.columns(2)
with col1:
    st.download_button(
        label="Download as .txt",
        data=transcript.encode("utf-8"),
        file_name=f"{base_name}.txt",
        mime="text/plain",
    )
with col2:
    st.download_button(
        label="Download as .docx",
        data=to_docx(transcript, base_name),
        file_name=f"{base_name}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
