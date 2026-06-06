import os
from typing import Callable, Optional
import db
from services import drive, openai_service, email, pdf


def process_assessment(
    assessment_id: int,
    video_path: str,
    progress_fn: Optional[Callable[[float, str], None]] = None,
) -> None:
    """
    Run the initial processing pipeline after student submission.
    PDF is NOT generated here — it is generated after mentor feedback is submitted.
    """
    def _progress(fraction: float, message: str) -> None:
        if progress_fn:
            progress_fn(fraction, message)

    db.update_assessment(assessment_id, status="processing")
    try:
        assessment = db.get_assessment_by_id(assessment_id)
        student_row = db.get_student_by_id(assessment["student_id"])
        mentor_row = db.get_mentor_by_id(student_row["mentor_id"])
        base_name = f"Mentoring Round {assessment['round']}. {student_row['name']}"

        _progress(0.0, "Uploading...")
        folder_id, folder_url = drive.create_student_round_folder(
            student_row["name"], assessment["round"], student_email=student_row.get("email") or ""
        )
        video_drive_url = drive.upload_file(
            video_path, folder_id, f"{base_name}. Recording.mp4"
        )
        db.update_assessment(
            assessment_id,
            drive_folder_url=folder_url,
            drive_folder_id=folder_id,
            video_drive_url=video_drive_url,
        )

        _progress(0.25, "Uploading...")
        audio_path = video_path.rsplit(".", 1)[0] + ".mp3"
        openai_service.extract_audio(video_path, audio_path)

        _progress(0.40, "Uploading...")
        transcript = openai_service.transcribe(audio_path)

        for path in (audio_path, video_path):
            try:
                os.remove(path)
            except FileNotFoundError:
                pass

        _progress(0.50, "Uploading...")
        from docx import Document
        from io import BytesIO
        transcript_path = f"uploads/transcript_{assessment_id}.docx"
        os.makedirs("uploads", exist_ok=True)
        doc = Document()
        doc.add_heading(f"{base_name}. Transcript", 0)
        for line in transcript.split("\n"):
            doc.add_paragraph(line)
        doc.save(transcript_path)
        drive.upload_file(transcript_path, folder_id, f"{base_name}. Transcript.docx")
        try:
            os.remove(transcript_path)
        except FileNotFoundError:
            pass

        db.update_assessment(assessment_id, transcript=transcript)

        _progress(0.65, "Uploading...")
        ai_review_content = openai_service.generate_ai_review(assessment, transcript)
        db.save_ai_review(assessment_id, ai_review_content)

        _progress(0.90, "Uploading...")
        app_url = os.environ.get("APP_URL", "http://localhost:8501")
        transcript_url = f"{app_url}/Transcript?assessment_id={assessment_id}"
        ai_review_url = f"{app_url}/AI_Review?assessment_id={assessment_id}"
        mentor_review_url = f"{app_url}/Mentor_Review?assessment_id={assessment_id}"
        email.send_mentor_notification(
            mentor_email=mentor_row["email"],
            mentor_name=mentor_row["name"],
            mentor_id=mentor_row["id"],
            student_name=student_row["name"],
            round_num=assessment["round"],
            video_drive_url=video_drive_url,
            transcript_url=transcript_url,
            ai_review_url=ai_review_url,
            mentor_review_url=mentor_review_url,
        )

        if student_row.get("email"):
            email.send_student_confirmation(
                student_email=student_row["email"],
                student_name=student_row["name"],
                round_num=assessment["round"],
                drive_folder_url=folder_url,
            )

        db.update_assessment(assessment_id, status="complete", error_message=None)
        _progress(1.0, "Uploading...")

    except Exception as exc:
        db.update_assessment(assessment_id, status="error", error_message=str(exc))
        raise


def generate_and_send_pdf(assessment_id: int) -> str:
    """
    Generate the assessment PDF and email it to both mentor and coach.
    Called after mentor submits feedback. Returns the Drive URL of the PDF.
    """
    assessment = db.get_assessment_by_id(assessment_id)
    student_row = db.get_student_by_id(assessment["student_id"])
    mentor_row = db.get_mentor_by_id(student_row["mentor_id"])
    ai_review = db.get_ai_review(assessment_id)
    feedback = db.get_mentor_feedback(assessment_id)
    base_name = f"Mentoring Round {assessment['round']}. {student_row['name']}"

    pdf_path = f"uploads/assessment_{assessment_id}.pdf"
    os.makedirs("uploads", exist_ok=True)
    import json as _json
    mentor_ratings = _json.loads(feedback.get("mentor_ratings") or "{}") if feedback else {}
    pdf.generate_pdf(
        assessment=assessment,
        student_name=student_row["name"],
        mentor_name=mentor_row["name"] if mentor_row else "",
        transcript="",
        mentor_feedback=feedback["feedback_text"] if feedback else "",
        mentor_ratings=mentor_ratings,
        output_path=pdf_path,
    )

    folder_id = assessment.get("drive_folder_id")
    pdf_drive_url = drive.upload_file(pdf_path, folder_id, f"{base_name}. Assessment.pdf")
    try:
        os.remove(pdf_path)
    except FileNotFoundError:
        pass

    db.update_assessment(assessment_id, pdf_drive_url=pdf_drive_url)

    app_url = os.environ.get("APP_URL", "http://localhost:8501")
    email.send_completion_notification(
        mentor_email=mentor_row["email"],
        mentor_name=mentor_row["name"],
        student_email=student_row.get("email") or "",
        student_name=student_row["name"],
        round_num=assessment["round"],
        pdf_drive_url=pdf_drive_url,
    )

    return pdf_drive_url
